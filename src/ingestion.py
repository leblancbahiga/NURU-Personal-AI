#!/usr/bin/env python3
"""
ingestion.py — Parsing, chunking et indexation de documents pour NURU.

Gère :
  - Parsing de PDF (PyMuPDF), DOCX (python-docx), TXT/MD
  - Chunking sémantique (par paragraphe/section) avec overlap
  - Empreinte MD5 pour éviter la ré-indexation
  - Suspension intelligente si RAM < 1.5 Go ou batterie < 20%
"""

import hashlib
import os
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None


# ── Types ──
@dataclass
class Document:
    """Document parsé et chunké."""
    filepath: str
    filename: str
    file_hash: str          # MD5 du fichier complet
    mime_type: str
    total_chunks: int
    chunks: list["Chunk"] = field(default_factory=list)
    indexed_at: float = field(default_factory=time.time)


@dataclass
class Chunk:
    """Segment de document."""
    text: str
    chunk_index: int
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    token_estimate: int = 0


SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
}


def get_mime(path: str) -> str:
    """Détecte le type MIME depuis l'extension."""
    return SUPPORTED_EXTENSIONS.get(Path(path).suffix.lower(), "text/plain")


# ── Empreinte MD5 ──

def compute_md5(filepath: str) -> str:
    """Calcule l'empreinte MD5 d'un fichier."""
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        # Lire en chunks de 64 Ko pour les gros fichiers
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# ── Parsing ──

def parse_pdf(filepath: str) -> tuple[str, list[int]]:
    """
    Extrait le texte d'un PDF.
    Retourne (texte_complet, [numéros_de_page_par_paragraphe]).
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) non installé — impossible de parser les PDF")

    doc = fitz.open(filepath)
    text_parts = []
    page_map = []

    for page_num, page in enumerate(doc, 1):
        blocks = page.get_text("blocks")
        for block in blocks:
            text = block[4].strip() if len(block) > 4 else ""
            if text:
                text_parts.append(text)
                page_map.append(page_num)

    doc.close()
    return "\n\n".join(text_parts), page_map


def parse_docx(filepath: str) -> str:
    """Extrait le texte d'un fichier DOCX."""
    if docx is None:
        raise ImportError("python-docx non installé — impossible de parser les DOCX")

    d = docx.Document(filepath)
    paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def parse_text(filepath: str) -> str:
    """Extrait le texte d'un fichier texte simple."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_document(filepath: str) -> tuple[str, str, list[int] | None]:
    """
    Parse un document selon son type.
    Retourne (texte, mime_type, page_map_optionnel).
    """
    ext = Path(filepath).suffix.lower()
    mime = get_mime(filepath)

    if ext == ".pdf":
        text, page_map = parse_pdf(filepath)
        return text, mime, page_map
    elif ext == ".docx":
        text = parse_docx(filepath)
        return text, mime, None
    else:
        # TXT, MD, CSV, etc.
        text = parse_text(filepath)
        return text, mime, None


# ── Chunking sémantique ──

def estimate_tokens(text: str) -> int:
    """Estimation rapide du nombre de tokens (~4 caractères par token en français)."""
    return len(text) // 4


def semantic_chunk(text: str, page_map: list[int] | None = None,
                   chunk_size: int = 512, overlap: int = 64) -> list[Chunk]:
    """
    Chunking sémantique : découpe par paragraphe/section, puis fusionne
    jusqu'à atteindre chunk_size tokens sans dépasser.

    Args:
        text: Texte complet du document
        page_map: Liste des numéros de page par paragraphe (PDF)
        chunk_size: Taille cible en tokens (défaut: 512)
        overlap: Chevauchement en tokens (défaut: 64)

    Retourne:
        Liste de Chunk
    """
    # Découpage en paragraphes
    paragraphs = re.split(r"\n\s*\n", text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks = []
    current_chunk = []
    current_tokens = 0
    chunk_idx = 0
    para_start_idx = 0  # Index du premier paragraphe dans ce chunk

    for i, para in enumerate(paragraphs):
        para_tokens = estimate_tokens(para)

        # Si un seul paragraphe dépasse la taille max, on le découpe en phrases
        if para_tokens > chunk_size and not current_chunk:
            # Découpage en phrases
            sentences = re.split(r"(?<=[.!?])\s+", para)
            for sent in sentences:
                sent_tokens = estimate_tokens(sent)
                if current_tokens + sent_tokens > chunk_size and current_chunk:
                    # Finaliser le chunk actuel
                    chunk_text = " ".join(current_chunk)
                    page = page_map[para_start_idx] if page_map and para_start_idx < len(page_map) else None
                    chunks.append(Chunk(
                        text=chunk_text,
                        chunk_index=chunk_idx,
                        page_number=page,
                        token_estimate=current_tokens,
                    ))
                    chunk_idx += 1

                    # Overlap : garder les dernières phrases
                    overlap_text = current_chunk
                    current_chunk = []
                    current_tokens = 0
                    for old_sent in reversed(overlap_text):
                        old_tokens = estimate_tokens(old_sent)
                        if current_tokens + old_tokens > overlap:
                            break
                        current_chunk.insert(0, old_sent)
                        current_tokens += old_tokens
                    para_start_idx = i

                current_chunk.append(sent)
                current_tokens += sent_tokens

            # Finaliser le dernier morceau
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                page = page_map[para_start_idx] if page_map and para_start_idx < len(page_map) else None
                chunks.append(Chunk(
                    text=chunk_text,
                    chunk_index=chunk_idx,
                    page_number=page,
                    token_estimate=current_tokens,
                ))
                chunk_idx += 1
                current_chunk = []
                current_tokens = 0
            continue

        # Ajout normal du paragraphe
        if current_tokens + para_tokens > chunk_size and current_chunk:
            # Finaliser le chunk actuel
            chunk_text = "\n\n".join(current_chunk)
            page = page_map[para_start_idx] if page_map and para_start_idx < len(page_map) else None
            chunks.append(Chunk(
                text=chunk_text,
                chunk_index=chunk_idx,
                page_number=page,
                token_estimate=current_tokens,
            ))
            chunk_idx += 1

            # Overlap : garder les derniers paragraphes
            overlap_paras = current_chunk
            current_chunk = []
            current_tokens = 0
            for old_para in reversed(overlap_paras):
                old_tokens = estimate_tokens(old_para)
                if current_tokens + old_tokens > overlap:
                    break
                current_chunk.insert(0, old_para)
                current_tokens += old_tokens
            para_start_idx = i - len(current_chunk) + len(overlap_paras) - len(current_chunk)

        current_chunk.append(para)
        current_tokens += para_tokens

    # Dernier chunk
    if current_chunk:
        chunk_text = "\n\n".join(current_chunk)
        page = page_map[para_start_idx] if page_map and para_start_idx < len(page_map) else None
        chunks.append(Chunk(
            text=chunk_text,
            chunk_index=chunk_idx,
            page_number=page,
            token_estimate=current_tokens,
        ))

    return chunks


# ── Gestionnaire de ressources ──

def check_system_resources() -> tuple[bool, str]:
    """
    Vérifie si le système a assez de ressources pour l'indexation.
    Adapté pour macOS (Apple Silicon) — utilise la pression mémoire
    plutôt que les pages libres, car macOS met en cache les fichiers.
    Retourne (ok, message).
    """
    # Vérification pression mémoire (macOS)
    try:
        import subprocess
        # memory_pressure donne le ratio de pression
        result = subprocess.run(
            ["memory_pressure"],
            capture_output=True, text=True, timeout=5
        )
        for line in result.stdout.split("\n"):
            if "pressure" in line.lower() and "percent" in line.lower():
                # Ex: "System-wide memory pressure percentage: 65%"
                try:
                    pct = int(''.join(c for c in line if c.isdigit() or c == '%').rstrip('%'))
                    if pct > 90:
                        return False, f"Pression mémoire trop élevée ({pct}%)"
                except (ValueError, IndexError):
                    pass
    except Exception:
        pass  # Ne pas bloquer

    # Vérification batterie
    try:
        result = subprocess.run(
            ["pmset", "-g", "batt"],
            capture_output=True, text=True, timeout=5
        )
        if "Battery Power" in result.stdout:
            for line in result.stdout.split("\n"):
                if "%" in line:
                    try:
                        pct = int(line.split("\t")[-1].split("%")[0])
                        if pct < 20:
                            return False, f"Batterie trop faible ({pct}% < 20%)"
                    except (ValueError, IndexError):
                        pass
    except Exception:
        pass

    return True, "OK"


def index_document(filepath: str, chunk_size: int = 512, overlap: int = 64) -> Optional[Document]:
    """
    Parse, chunke et prépare un document pour l'indexation vectorielle.

    Args:
        filepath: Chemin du fichier
        chunk_size: Taille des chunks en tokens
        overlap: Chevauchement en tokens

    Retourne:
        Document avec chunks, ou None si échec
    """
    path = Path(filepath)

    if not path.exists():
        print(f"✗ Fichier introuvable : {filepath}")
        return None

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        print(f"  ⚠ Extension non supportée : {ext}")
        return None

    try:
        # Vérification des ressources
        ok, msg = check_system_resources()
        if not ok:
            print(f"  ⚠ Suspension : {msg}")
            return None

        # MD5
        file_hash = compute_md5(filepath)

        # Parsing
        print(f"  📄 Parsing : {path.name}...", end=" ", flush=True)
        text, mime, page_map = parse_document(str(path))

        if not text.strip():
            print("(vide)")
            return None

        print(f"({len(text)} chars)", end=" ", flush=True)

        # Chunking
        chunks = semantic_chunk(text, page_map, chunk_size, overlap)
        print(f"→ {len(chunks)} chunks")

        return Document(
            filepath=str(path),
            filename=path.name,
            file_hash=file_hash,
            mime_type=mime,
            total_chunks=len(chunks),
            chunks=chunks,
        )

    except Exception as e:
        print(f"  ✗ Erreur : {e}")
        return None
